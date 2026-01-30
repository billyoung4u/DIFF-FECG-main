from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_guide_doc():
    doc = Document()

    # --- 标题 ---
    heading = doc.add_heading('AI 算法服务接入与集成操作指南', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('版本: V1.0 | 适用对象: 系统开发/Java 工程师')

    # --- 1. 任务概述 ---
    doc.add_heading('1. 任务概述', level=1)
    p = doc.add_paragraph(
        '本项目的核心是将封装好的“胎儿心电 AI 算法”集成到现有业务系统中。'
        'AI 服务已封装为 Docker 镜像，独立运行。'
    )

    doc.add_paragraph('业务系统的职责仅包括以下三点：', style='List Number')
    doc.add_paragraph('部署服务：在服务器启动 Docker 容器。', style='List Number')
    doc.add_paragraph('发送数据：通过 HTTP 接口上传原始 ECG 数据。', style='List Number')
    doc.add_paragraph('结果展示：解析返回的 JSON 数据并绘制波形。', style='List Number')

    # --- 2. 第一阶段 ---
    doc.add_heading('2. 第一阶段：部署篇（环境准备）', level=1)

    doc.add_heading('2.1 安装 Docker', level=2)
    doc.add_paragraph('请在目标服务器（或测试机）上安装 Docker 环境。验证方法：执行 docker -v。')

    doc.add_heading('2.2 获取镜像包', level=2)
    doc.add_paragraph('请接收算法团队提供的项目压缩包 project.zip 并解压。')

    doc.add_heading('2.3 启动服务', level=2)
    doc.add_paragraph('在解压后的项目根目录下，依次执行以下两条命令：')

    # 代码块样式
    code_build = "docker build -f deployment/Dockerfile -t fecg-ai-service:v1 ."
    code_run = "docker run -d -p 8000:8000 fecg-ai-service:v1"

    p = doc.add_paragraph()
    runner = p.add_run(f"1. 构建镜像：\n{code_build}\n\n2. 启动容器：\n{code_run}")
    runner.font.name = 'Courier New'
    runner.font.color.rgb = RGBColor(0, 51, 102)  # 深蓝色代码

    doc.add_paragraph('验证服务状态：访问 http://localhost:8000/docs')

    # --- 3. 第二阶段 ---
    doc.add_heading('3. 第二阶段：开发篇（Java 代码接入）', level=1)

    doc.add_heading('3.1 引入依赖', level=2)
    doc.add_paragraph('推荐使用 Hutool 工具包：')
    maven_code = """<dependency>
    <groupId>cn.hutool</groupId>
    <artifactId>hutool-all</artifactId>
    <version>5.8.16</version>
</dependency>"""
    p = doc.add_paragraph()
    runner = p.add_run(maven_code)
    runner.font.name = 'Courier New'
    runner.font.size = Pt(9)

    doc.add_heading('3.2 Java 调用代码示例', level=2)
    doc.add_paragraph('类名：EcgAiClient.java')

    java_code = """
import cn.hutool.http.HttpRequest;
import cn.hutool.json.JSONObject;
import cn.hutool.json.JSONUtil;

public class EcgAiClient {
    private static final String API_URL = "http://127.0.0.1:8000/predict";

    public static void main(String[] args) {
        // 1. 准备数据 (500Hz)
        double[] rawData = {0.012, 0.034, 0.15, -0.02, 0.0, 0.012}; 

        // 2. 组装参数
        JSONObject requestBody = new JSONObject();
        requestBody.set("patient_id", "TEST_001");
        requestBody.set("fs", 500);
        requestBody.set("data", rawData);

        try {
            // 3. 发送 POST 请求
            String response = HttpRequest.post(API_URL)
                    .body(requestBody.toString())
                    .timeout(10000)
                    .execute().body();

            // 4. 解析结果
            JSONObject jsonRes = JSONUtil.parseObj(response);
            if (jsonRes.getInt("code") == 200) {
                JSONObject result = jsonRes.getJSONObject("result");
                Double[] fecgSignal = result.getBeanList("fecg_signal", Double.class).toArray(new Double[0]);
                double fhr = result.getDouble("fhr_bpm");

                System.out.println("胎心率: " + fhr);
                // TODO: 绘图
            }
        } catch (Exception e) {
            e.printStackTrace();
        }
    }
}
"""
    p = doc.add_paragraph()
    runner = p.add_run(java_code)
    runner.font.name = 'Courier New'
    runner.font.size = Pt(8)

    # --- 4. 第三阶段 ---
    doc.add_heading('4. 第三阶段：展示篇（前端绘图）', level=1)
    doc.add_paragraph('前端收到 Java 后端转发的 fecg_signal 数组后，需进行波形渲染。')
    doc.add_paragraph('推荐使用 ECharts 或 JavaFX Charts。', style='List Bullet')
    doc.add_paragraph('X轴为时间，Y轴为 fecg_signal 数值。', style='List Bullet')

    # --- 5. 职责清单 ---
    doc.add_heading('5. 职责清单', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '角色'
    hdr_cells[1].text = '需要做的事'
    hdr_cells[2].text = '不需要做的事'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Java 开发'
    row_cells[1].text = '1. 装 Docker\n2. 发 HTTP 请求\n3. 画图'
    row_cells[2].text = '1. 不用装 Python\n2. 不用懂模型'

    # 保存文件
    file_name = 'AI_Service_Integration_Guide.docx'
    doc.save(file_name)
    print(f"✅ 文档已生成：{file_name}")


if __name__ == "__main__":
    create_guide_doc()